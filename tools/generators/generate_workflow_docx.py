"""
Generate a comprehensive Word document summarizing the RL Trading System workflow.
Output: RL_Trading_Workflow.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_para(doc, text, bold=False, italic=False, size=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    # add light gray background via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p


def add_table(doc, headers, rows, header_bg='2E5BBA'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    # Header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], header_bg)
    # Body
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
    return table


# ============================================================
# CREATE DOCUMENT
# ============================================================
doc = Document()

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ----------------- TITLE PAGE -----------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('\n\n\n\nRL Trading System')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x2E, 0x5B, 0xBA)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Complete Workflow Summary')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('PPO Reinforcement Learning โ ONNX โ MetaTrader 5')
run.italic = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph('\n\n\n')

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('EUR/USD Trading Bot\n').bold = True
info.add_run('Python (PyCaret Trainer GUI) + MT5 Expert Advisor\n')
info.add_run('Version V11 (with 10 Candle Patterns)\n')
info.add_run('Date: 2026-05-09')

doc.add_page_break()

# ----------------- TABLE OF CONTENTS -----------------
add_heading(doc, 'Table of Contents', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))
toc_items = [
    '1. System Overview',
    '2. Architecture Diagram',
    '3. Phase 1: Data Collection (MT5)',
    '4. Phase 2: Data Preparation',
    '5. Phase 3: Feature Engineering',
    '6. Phase 4: Training (PPO V4 Reward)',
    '7. Phase 5: Backtest & Validation',
    '8. Phase 6: MT5 Deployment (ONNX)',
    '9. Phase 7: Iterate & Improve',
    '10. Versions Journey (V2 โ V11)',
    '11. Tools Reference',
    '12. Quick Start Guide',
    '13. File Structure',
    '14. Troubleshooting',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Inches(0.3)

doc.add_page_break()

# ----------------- 1. SYSTEM OVERVIEW -----------------
add_heading(doc, '1. System Overview', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))
add_para(doc,
    'ระบบนี้คือ end-to-end pipeline สำหรับสร้าง RL Trading Bot ที่เทรน '
    'ด้วย PPO (Proximal Policy Optimization) บน Python แล้วนำไป deploy '
    'บน MetaTrader 5 ผ่าน ONNX format โดยมี GUI ภาษาไทย/อังกฤษ '
    'ครอบคลุมทุกขั้นตอนตั้งแต่เก็บข้อมูล โหลด feature engineering '
    'เทรน backtest และ export EA สำหรับใช้งานจริง'
)

add_heading(doc, 'Key Components', level=2)
add_bullet(doc, 'GUI Trainer (rl_app.py): CustomTkinter app, ~3500 บรรทัด, branding ปรับเองได้')
add_bullet(doc, 'Trading Environment (trading_env.py): Custom Gymnasium env + V4 reward')
add_bullet(doc, 'PPO Training (stable-baselines3 2.x)')
add_bullet(doc, 'ONNX Export (export_to_onnx.py): generates .onnx + _config.mqh + _EA.mq5')
add_bullet(doc, 'MT5 EA Template (ML_RL_Trader_template.mq5): #resource embed, dynamic feature mapping')
add_bullet(doc, 'CandlePatterns Indicator (CandlePatterns.mq5): 10 patterns + tunable thresholds')
add_bullet(doc, 'DataCollector EA (DataCollector_v4.mq5): exports CSV with 67 columns')

doc.add_page_break()

# ----------------- 2. ARCHITECTURE -----------------
add_heading(doc, '2. Architecture Diagram', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))
add_code(doc,
    'โโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโ\n'
    'โ                MT5 (DataCollector_v4)                 โ\n'
    'โ        Export EUR/USD H1 + 67 columns CSV             โ\n'
    'โโโโโโโโโโโโโโโโโโโโโโโโโโโฌโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโ\n'
    '                          โ\n'
    '                          โผ\n'
    'โโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโ\n'
    'โ            PyCaret Trainer GUI (rl_app.py)            โ\n'
    'โ  Tools: Import โ Feature Eng โ Train โ Backtest      โ\n'
    'โ          โ Export ONNX โ Deploy MT5 EA               โ\n'
    'โโโโโโโโโโโโโโโโโโโโโโโโโโโฌโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโ\n'
    '                          โ\n'
    '                          โผ\n'
    'โโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโ\n'
    'โ              MT5 EA (ML_RL_Trader_*.ex5)              โ\n'
    'โ   ONNX embedded (#resource) + iCustom indicators      โ\n'
    'โ        + Session filter + Risk management             โ\n'
    'โโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโโ\n'
)

doc.add_page_break()

# ----------------- 3. PHASE 1 -----------------
add_heading(doc, '3. Phase 1: Data Collection (MT5)', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))
add_heading(doc, 'Goal', level=2)
add_para(doc, 'รวบรวมข้อมูล historical EUR/USD H1 พร้อม indicators และ candle patterns ออกมาเป็น CSV')

add_heading(doc, 'Tool', level=2)
add_para(doc, 'mt5_files/MQL5/Experts/DataCollector_v4.mq5', bold=True)

add_heading(doc, 'Output Columns (67 total)', level=2)
add_bullet(doc, 'Time + OHLCV (6 cols)')
add_bullet(doc, 'V3 indicators: RSI, ATR, ADX, MACD, BB, Stoch, ฯลฯ (~30 cols)')
add_bullet(doc, 'Phase A multi-timeframe: d1_rsi, d1_trend, d1_atr (correlation 0.138 กับ future_return)')
add_bullet(doc, 'Candle patterns (10 cols): Hammer, Engulfing, Inside, Outside, Star, Soldiers, Marubozu, Harami, Piercing, MatHold')
add_bullet(doc, 'future_return (target สำหรับ analysis)')

add_heading(doc, 'Steps', level=2)
add_para(doc, '1. Copy CandlePatterns.mq5 ไปที่ MQL5/Indicators/, compile')
add_para(doc, '2. Copy DataCollector_v4.mq5 ไปที่ MQL5/Experts/, compile')
add_para(doc, '3. Attach EA บน chart EUR/USD H1, ตั้ง start/end date')
add_para(doc, '4. รันแล้ว CSV จะออกที่ MQL5/Files/')

doc.add_page_break()

# ----------------- 4. PHASE 2 -----------------
add_heading(doc, '4. Phase 2: Data Preparation', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))
add_heading(doc, 'Goal', level=2)
add_para(doc, 'Import CSV เข้า GUI, แยก train/test, normalize')

add_heading(doc, 'Tool', level=2)
add_para(doc, 'rl_app.py โ Tools page โ Card 1 (Import Data)', bold=True)

add_heading(doc, 'Features', level=2)
add_bullet(doc, 'Auto-detect v3 (57 cols) vs v4 (67 cols)')
add_bullet(doc, 'แสดงสถิติ + missing values')
add_bullet(doc, 'Normalize (z-score) แยกตาม train set')
add_bullet(doc, 'แยก train/test ตามวันที่ (default: 80/20)')

add_heading(doc, 'Output', level=2)
add_code(doc,
    'data/eurusd_h1_train_norm.csv\n'
    'data/eurusd_h1_test_norm.csv\n'
    'data/normalization_stats.json\n'
)

doc.add_page_break()

# ----------------- 5. PHASE 3 -----------------
add_heading(doc, '5. Phase 3: Feature Engineering', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))
add_heading(doc, 'Goal', level=2)
add_para(doc, 'วิเคราะห์ correlation feature กับ future_return เลือก feature เด่น')

add_heading(doc, 'Tool', level=2)
add_para(doc, 'rl_app.py โ Tools page โ Card 4 (Feature Analysis)', bold=True)

add_heading(doc, 'Phase A Discovery', level=2)
add_para(doc, 'd1_rsi (Daily RSI) มี correlation 0.138 กับ future_return ', bold=True)
add_para(doc, 'แรงกว่า indicator มาตรฐานถึง 3.4 เท่า')

add_heading(doc, 'Top Features (จาก analysis)', level=2)
features = [
    ['d1_rsi', '0.138', 'Multi-timeframe RSI'],
    ['d1_trend', '0.095', 'Daily trend direction'],
    ['rsi', '0.041', 'Hourly RSI'],
    ['adx', '0.038', 'Trend strength'],
    ['macd_hist', '0.029', 'MACD histogram'],
    ['candle_engulfing', '0.024', 'Bullish/Bearish engulfing'],
    ['candle_mat_hold', '0.018', 'Continuation pattern'],
]
add_table(doc, ['Feature', 'Correlation', 'Note'], features)

doc.add_page_break()

# ----------------- 6. PHASE 4: TRAINING -----------------
add_heading(doc, '6. Phase 4: Training (PPO V4 Reward)', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))
add_heading(doc, 'Goal', level=2)
add_para(doc, 'เทรน PPO agent ให้ตัดสินใจ Buy / Hold / Sell บน EUR/USD H1')

add_heading(doc, 'Tool', level=2)
add_para(doc, 'rl_app.py โ Tools page โ Card 2 (Train RL)', bold=True)

add_heading(doc, 'Custom Environment (trading_env.py)', level=2)
add_bullet(doc, 'Observation: window=20 candles ร feature_count')
add_bullet(doc, 'Action: 0=Hold, 1=Buy, 2=Sell')
add_bullet(doc, 'Spread baked in entry/exit price (ตรงกับ live backtest)')
add_bullet(doc, 'Random episode start: t = random[W-1, max_start] (95.7% data coverage)')

add_heading(doc, 'V4 Reward Function', level=2)
add_code(doc,
    '# Trade closed PnL\n'
    'reward += trade_closed_pnl * 50.0\n'
    '\n'
    '# Threshold profit bonus (กระตุ้นให้ปิดกำไร)\n'
    'if trade_closed_pnl > 0.005:  # > 0.5%\n'
    '    reward += 1.0\n'
    '\n'
    '# Trade penalty (ป้องกัน overtrading)\n'
    'if action != 0:\n'
    '    reward -= 0.005\n'
)

add_heading(doc, 'Hyperparameters', level=2)
hp = [
    ['learning_rate', '3e-4'],
    ['n_steps', '2048'],
    ['batch_size', '64'],
    ['n_epochs', '10'],
    ['gamma', '0.99'],
    ['gae_lambda', '0.95'],
    ['clip_range', '0.2'],
    ['total_timesteps', '500,000 - 2,000,000'],
]
add_table(doc, ['Parameter', 'Value'], hp)

doc.add_page_break()

# ----------------- 7. PHASE 5 -----------------
add_heading(doc, '7. Phase 5: Backtest & Validation', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))
add_heading(doc, 'Goal', level=2)
add_para(doc, 'ทดสอบ model บน test set + walk-forward เพื่อยืนยันว่าไม่ overfit')

add_heading(doc, 'Tool', level=2)
add_para(doc, 'rl_app.py โ Tools page โ Card 3 (Backtest)', bold=True)

add_heading(doc, 'Metrics', level=2)
add_bullet(doc, 'Profit Factor (PF) - target > 1.0')
add_bullet(doc, 'Win Rate %')
add_bullet(doc, 'Max Drawdown')
add_bullet(doc, 'Sharpe Ratio')
add_bullet(doc, 'จำนวน trades')
add_bullet(doc, 'Equity curve (interactive Plotly chart)')

add_heading(doc, 'Confidence Threshold', level=2)
add_para(doc, 'ใช้ argmax + softmax confidence filter (default 0.95)')
add_para(doc, 'เพื่อให้ EA trade เฉพาะตอนที่ model มั่นใจสูง โดย V10+ confidence 0.95 ทำให้ PF > 1.0 ได้')

doc.add_page_break()

# ----------------- 8. PHASE 6 -----------------
add_heading(doc, '8. Phase 6: MT5 Deployment (ONNX)', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))
add_heading(doc, 'Goal', level=2)
add_para(doc, 'Convert PPO โ ONNX โ embed ใน MT5 EA สำหรับเทรดจริง')

add_heading(doc, 'Tool', level=2)
add_para(doc, 'rl_app.py โ Tools page โ Card 5 (Export to ONNX)', bold=True)

add_heading(doc, 'Process', level=2)
add_para(doc, '1. PolicyWrapper หุ้ม actor ของ PPO + softmax')
add_para(doc, '2. torch.onnx.export() (PyTorch dynamo exporter, ไม่มี dynamic_axes)')
add_para(doc, '3. Consolidate external data: onnx.save_model(save_as_external_data=False)')
add_para(doc, '4. Generate _config.mqh (feature list + spread + confidence)')
add_para(doc, '5. Generate _EA.mq5 จาก template (replace placeholders)')

add_heading(doc, 'Output Files', level=2)
add_code(doc,
    '<model_name>.onnx              # ONNX model\n'
    '<model_name>_config.mqh        # Feature list + settings\n'
    '<model_name>_EA.mq5            # Generated EA\n'
)

add_heading(doc, 'EA Architecture', level=2)
add_bullet(doc, '#resource "\\\\Files\\\\<model>.onnx" - embed ONNX ใน .ex5')
add_bullet(doc, 'OnnxCreateFromBuffer() - load จาก resource (ใช้กับ Tester ได้)')
add_bullet(doc, 'iCustom("CandlePatterns", ...) - ถ้า model ใช้ candle features')
add_bullet(doc, 'RL_BuildFeatureMap() - dynamic feature mapping (รองรับ V10/V11/อื่นๆ)')
add_bullet(doc, 'IsAllowedSession() - SymbolInfoSessionTrade (broker-aware)')
add_bullet(doc, 'Session filter ใช้กับทุก action รวมถึง Close')

add_heading(doc, 'Dynamic Feature Mapping', level=2)
add_para(doc, 'แทนที่จะ hard-code feature count, EA ใช้ master list 75 features')
add_para(doc, 'แล้ว RL_BuildFeatureMap() build index map จาก RL_FEATURE_NAMES')
add_para(doc, 'ทำให้ EA เดียวกัน support model ต่าง version (V10=58, V11=58, รุ่นใหม่ๆ)')

doc.add_page_break()

# ----------------- 9. PHASE 7 -----------------
add_heading(doc, '9. Phase 7: Iterate & Improve', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))
add_para(doc, 'Loop กลับไป Phase 3 ถ้า:')
add_bullet(doc, 'PF < 1.0 บน test set')
add_bullet(doc, 'Drawdown สูงเกินรับได้')
add_bullet(doc, 'Win rate ต่ำผิดปกติ')
add_para(doc, '\nแนวทางปรับปรุง:')
add_bullet(doc, 'เพิ่ม feature ใหม่ (เช่น session indicator, news flag)')
add_bullet(doc, 'ปรับ V4 reward weights')
add_bullet(doc, 'เพิ่ม timesteps (1M โ 2M)')
add_bullet(doc, 'ปรับ confidence threshold')
add_bullet(doc, 'ลอง walk-forward retraining')

doc.add_page_break()

# ----------------- 10. VERSIONS JOURNEY -----------------
add_heading(doc, '10. Versions Journey (V2 โ V11)', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))
versions = [
    ['V2', 'Baseline reward', 'PF 0.6', 'Reward hacking - hold มากเกินไป'],
    ['V3', 'Penalty for hold', 'PF 0.7', 'เริ่มเทรดบ้าง แต่ขาดทุน'],
    ['V4', '*50 multiplier + threshold bonus', 'PF 0.85', 'ดีขึ้นมาก แต่ยังไม่ profit'],
    ['V5-V8', 'ทดลอง feature ต่างๆ', 'PF 0.7-0.9', 'ค่อยๆ ปรับ'],
    ['V9', 'Phase A multi-timeframe', 'PF 0.95', 'ใกล้ break-even'],
    ['V10', 'Confidence 0.95 + Phase A', 'PF 1.03', 'ตัวแรกที่ profit!'],
    ['V11', 'V10 + 10 candle patterns', 'Testing', 'ต้องใช้ EA ใหม่ที่ support candles'],
]
add_table(doc, ['Version', 'Key Change', 'Result', 'Note'], versions)

add_heading(doc, 'Critical Bug Discoveries', level=2)
add_bullet(doc, 'Random episode start bug: ใช้ t=W ตายตัว (1.8% data) โ fix random[W-1, max_start] (95.7%)')
add_bullet(doc, 'Phase 1 fix: window timing ตอน reset')
add_bullet(doc, 'Phase 2 fix: equity tracking with spread')
add_bullet(doc, 'V11 array out of range: feature count mismatch โ dynamic mapping')

doc.add_page_break()

# ----------------- 11. TOOLS REFERENCE -----------------
add_heading(doc, '11. Tools Reference', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))
tools = [
    ['Card 1', 'Import Data', 'Auto-detect v3/v4, normalize, train/test split'],
    ['Card 2', 'Train RL', 'PPO training with V4 reward + tensorboard logs'],
    ['Card 3', 'Backtest', 'Test set + walk-forward + Plotly equity chart'],
    ['Card 4', 'Feature Analysis', 'Correlation matrix + top features ranking'],
    ['Card 5', 'Export to ONNX', 'PPO โ ONNX + auto-generate _config.mqh + _EA.mq5'],
    ['Settings', 'Branding', 'Color presets + native color picker + JSON persistence'],
]
add_table(doc, ['Location', 'Tool', 'Function'], tools)

add_heading(doc, 'Branding Customization', level=2)
add_bullet(doc, 'Settings page โ Branding Customization card')
add_bullet(doc, 'Logo: ดึงจากไฟล์ได้ (PNG/JPG)')
add_bullet(doc, 'Color presets: Default, Warm, Ocean, Pure Black, Metafxclub')
add_bullet(doc, 'Native color picker: คลิก swatch เพื่อเปิด tkinter colorchooser')
add_bullet(doc, 'Persist: บันทึกใน branding_config.json (gitignored)')

doc.add_page_break()

# ----------------- 12. QUICK START -----------------
add_heading(doc, '12. Quick Start Guide', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))

add_heading(doc, 'Setup', level=2)
add_code(doc,
    '# 1. Install dependencies\n'
    '.venv\\Scripts\\python.exe -m pip install -r requirements.txt\n'
    '\n'
    '# 2. Run GUI\n'
    '.venv\\Scripts\\python.exe rl_app.py\n'
)

add_heading(doc, 'Train New Model (V11+)', level=2)
add_para(doc, '1. Open GUI โ Tools page')
add_para(doc, '2. Card 1: Import data/eurusd_h1_train_norm.csv (v4 67 cols)')
add_para(doc, '3. Card 4: Feature Analysis (เลือก features ที่ correlation สูง)')
add_para(doc, '4. Card 2: Train RL (set total_timesteps=1M, V4 reward)')
add_para(doc, '5. Card 3: Backtest (ตรวจ PF > 1.0)')
add_para(doc, '6. Card 5: Export to ONNX โ ได้ <name>.onnx + _config.mqh + _EA.mq5')

add_heading(doc, 'Deploy to MT5', level=2)
add_para(doc, '1. Copy mt5_files/MQL5/Indicators/CandlePatterns.mq5 โ MT5/MQL5/Indicators/')
add_para(doc, '2. Copy <name>.onnx โ MT5/MQL5/Files/')
add_para(doc, '3. Copy <name>_config.mqh โ MT5/MQL5/Include/')
add_para(doc, '4. Copy <name>_EA.mq5 โ MT5/MQL5/Experts/')
add_para(doc, '5. Compile ทั้งหมดใน MetaEditor')
add_para(doc, '6. Backtest ใน Strategy Tester (EUR/USD H1)')
add_para(doc, '7. ถ้า OK โ Attach EA บน live chart (demo first!)')

doc.add_page_break()

# ----------------- 13. FILE STRUCTURE -----------------
add_heading(doc, '13. File Structure', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))
add_code(doc,
    'pycaret_trainer/\n'
    'โโโ rl_app.py                         # GUI main (~3500 lines)\n'
    'โโโ trading_env.py                    # Custom Gymnasium env\n'
    'โโโ export_to_onnx.py                 # PPO โ ONNX converter\n'
    'โโโ backtest_live.py                  # Live backtest engine\n'
    'โโโ requirements.txt                  # Python deps\n'
    'โโโ branding_config.json              # User branding (gitignored)\n'
    'โโโ data/\n'
    'โ   โโโ eurusd_h1_train_norm.csv\n'
    'โ   โโโ eurusd_h1_test_norm.csv\n'
    'โ   โโโ normalization_stats.json\n'
    'โโโ models/\n'
    'โ   โโโ V10_*.zip                     # PPO checkpoints\n'
    'โ   โโโ V11_*.zip\n'
    'โ   โโโ V11_*.onnx                    # Exported ONNX\n'
    'โโโ mt5_files/MQL5/\n'
    '    โโโ Indicators/\n'
    '    โ   โโโ CandlePatterns.mq5        # 10 patterns indicator\n'
    '    โโโ Include/\n'
    '    โ   โโโ RL_Indicators.mqh         # Feature builder + dynamic map\n'
    '    โ   โโโ <model>_config.mqh        # Per-model config (generated)\n'
    '    โโโ Experts/\n'
    '        โโโ DataCollector_v4.mq5      # CSV exporter\n'
    '        โโโ ML_RL_Trader_template.mq5 # EA template\n'
    '        โโโ <model>_EA.mq5            # Generated per model\n'
)

doc.add_page_break()

# ----------------- 14. TROUBLESHOOTING -----------------
add_heading(doc, '14. Troubleshooting', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))
issues = [
    ['ModuleNotFoundError: onnxscript',
     'pip install onnxscript onnx onnxruntime'],
    ['ONNX error 5019 (MT5 Tester)',
     'ใช้ #resource directive + OnnxCreateFromBuffer'],
    ['ONNX split into .onnx + .onnx.data',
     'consolidate ด้วย onnx.save_model(save_as_external_data=False)'],
    ['MQH not found during compile',
     'ใช้ #include <...> แทน "..." สำหรับไฟล์ใน MQL5/Include/'],
    ['"Market closed" errors in Tester',
     'IsAllowedSession() with SymbolInfoSessionTrade() (broker-aware)'],
    ['"Array out of range" (V11)',
     'Feature count mismatch โ ใช้ EA template ที่มี dynamic feature mapping'],
    ['Blank GUI window',
     'state init + _build_ui() ต้องอยู่ใน __init__'],
    ['Branding ไม่ persist',
     'check branding_config.json อยู่ใน project root + ไม่ถูก gitignored จากการเขียน'],
]
add_table(doc, ['Problem', 'Solution'], issues)

# ----------------- FOOTER -----------------
doc.add_paragraph('\n\n')
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = end.add_run('โ End of Workflow Summary โ')
run.italic = True
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ----------------- SAVE -----------------
output_path = r'C:\Users\omesb\Documents\claude code\pycaret_trainer\RL_Trading_Workflow.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
