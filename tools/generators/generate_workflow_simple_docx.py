"""
Generate a beginner-friendly Word document for traders who don't know Python/RL.
Output: RL_Trading_Simple.docx

Style: ใช้ภาษาง่ายๆ + analogies จากโลกการเทรด, ไม่มี code, เน้นให้เห็นภาพรวม
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_para(doc, text, bold=False, italic=False, size=None, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    return p


def add_callout(doc, title, body, bg='FFF4CE', border='FFC107'):
    """กล่องคำเตือน/เคล็ดลับสีเหลือง"""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    cell.paragraphs[0].text = ''
    p1 = cell.paragraphs[0]
    run = p1.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x66, 0x00)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    return table


def add_analogy(doc, title, body):
    """กล่องเทียบเคียง (สีฟ้าอ่อน)"""
    return add_callout(doc, title, body, bg='E3F2FD', border='2196F3')


def add_table(doc, headers, rows, header_bg='2E5BBA'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], header_bg)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
    return table


# ============================================================
# CREATE DOCUMENT
# ============================================================
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(12)

# ----------------- TITLE PAGE -----------------
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run('\n\n\n')

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run('RL Trading Bot')
run.bold = True
run.font.size = Pt(40)
run.font.color.rgb = RGBColor(0x2E, 0x5B, 0xBA)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run('สอน AI ให้เทรดแทนเรา')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run('ฉบับเข้าใจง่าย สำหรับนักเทรดที่ไม่รู้จัก Python/AI')
run.italic = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph('\n\n')
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run('"เหมือนฝึกลูกศิษย์เทรดเดอร์ ให้กลายเป็น EA"\n\n')
r.italic = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('คู่ EUR/USD H1 | ใช้กับ MetaTrader 5\n')
info.add_run('Date: 2026-05-09')

doc.add_page_break()

# ----------------- บทนำ -----------------
add_heading(doc, 'อ่านก่อนเริ่ม', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))
add_para(doc,
    'เอกสารนี้เขียนสำหรับนักเทรดที่อยากรู้ว่า "RL Bot" คืออะไร และเขาฝึก AI '
    'ให้เทรดแทนเรายังไง โดยไม่ต้องรู้ Python ไม่ต้องรู้คณิตศาสตร์ '
    'ไม่ต้องเขียนโค้ด แค่อยากเข้าใจภาพรวมว่ามันทำงานยังไง '
    'แล้วผลที่ได้คือ EA ตัวหนึ่งที่ใส่ลง MT5 ได้เหมือนปกติ'
)

add_callout(doc,
    'TL;DR (สรุปสั้นมาก)',
    'เราเอาข้อมูลกราฟ EUR/USD ย้อนหลังหลายปี มาให้ AI เล่นซ้ำๆ เป็นล้านรอบ '
    'AI จะเรียนรู้เองว่า "ตอนไหนควรซื้อ ตอนไหนควรขาย ตอนไหนควรอยู่เฉย" '
    'พอ AI เก่งแล้ว เราก็แปลงสมองของมันเป็นไฟล์ ONNX แล้วฝังลง EA ของ MT5 '
    'ให้เทรดจริงในตลาด'
)

doc.add_page_break()

# ----------------- 1. RL คืออะไร -----------------
add_heading(doc, '1. RL คืออะไร? (อธิบายแบบบ้านๆ)', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))

add_para(doc, 'RL ย่อมาจาก ', bold=False)
add_para(doc, 'Reinforcement Learning = "การเรียนรู้แบบเสริมแรง"', bold=True, size=14, color=RGBColor(0x2E, 0x5B, 0xBA))

add_para(doc,
    'มันคือวิธีสอน AI ที่เลียนแบบวิธี "ฝึกสุนัข" หรือ "สอนเด็กขี่จักรยาน" '
    'คือไม่บอกคำตอบให้ตรงๆ แต่ให้ลองทำแล้วให้รางวัล/ลงโทษ '
    'ทำดี = ขนม, ทำพลาด = ดุ พอทำซ้ำเยอะๆ มันจะรู้เองว่าควรทำยังไง'
)

add_analogy(doc,
    '🐕 เปรียบเทียบ: ฝึกสุนัขนั่ง',
    'คุณไม่ได้สอนสุนัขด้วยตำราเรียน แต่พูดว่า "นั่ง" แล้วถ้ามันนั่งจริง = ให้ขนม '
    'ถ้าไม่ทำ = ไม่ได้ขนม ทำซ้ำ 100 ครั้ง สุนัขจะเชื่อมโยงเองว่า "นั่ง = ขนม" '
    'RL Bot ก็แบบเดียวกันเลย เพียงแต่เปลี่ยนจาก "นั่ง" เป็น "ซื้อ/ขาย/อยู่เฉย"'
)

add_heading(doc, 'องค์ประกอบของ RL ในแบบเทรดเดอร์', level=2)
elements = [
    ['Agent (ผู้เล่น)', 'AI เทรดเดอร์ที่เรากำลังฝึก'],
    ['Environment (สนาม)', 'กราฟ EUR/USD H1 ย้อนหลัง 5-10 ปี'],
    ['Action (สิ่งที่ทำได้)', 'ซื้อ / ขาย / อยู่เฉย (3 ตัวเลือก)'],
    ['Observation (สิ่งที่เห็น)', 'ราคา 20 แท่งล่าสุด + indicator + candle pattern'],
    ['Reward (รางวัล)', 'กำไร = บวก, ขาดทุน = ลบ, เทรดมั่ว = หักคะแนน'],
]
add_table(doc, ['ส่วนประกอบ', 'ในระบบของเรา'], elements)

doc.add_page_break()

# ----------------- 2. ทำไมไม่ใช้ if-else -----------------
add_heading(doc, '2. ทำไมไม่เขียน if-else แบบ EA ทั่วไป?', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))

add_para(doc, 'EA ปกติที่เราเจอกัน เขียนกฎไว้ในโค้ดประมาณนี้:')
add_para(doc, '   "ถ้า RSI < 30 → ซื้อ, ถ้า RSI > 70 → ขาย"', italic=True, color=RGBColor(0x66, 0x66, 0x66))

add_para(doc, '\nปัญหาของวิธีนี้คือ:')
add_bullet(doc, 'กฎตายตัว ตลาดเปลี่ยน กฎไม่เปลี่ยนตาม')
add_bullet(doc, 'คนเขียนต้อง "เดา" เอาเองว่ากฎไหนดี')
add_bullet(doc, 'ถ้ามี indicator 30 ตัว ใครจะรู้ว่า combination ไหนถูก?')
add_bullet(doc, 'ตลาดมีบริบท (context) เยอะมาก กฎ if-else จับไม่หมด')

add_heading(doc, 'RL ทำต่างยังไง', level=2)
add_para(doc,
    'RL ไม่ต้องให้เราบอกกฎเลยสักข้อ! เราแค่ให้มัน "เห็น" indicator ทั้งหมด '
    'แล้วบอกแค่ว่า "เป้าหมายคือทำกำไร" จากนั้น AI จะลองสุ่มซื้อ-ขายเอง '
    'รอบแรกๆ มันโง่มาก เทรดมั่วและขาดทุน แต่พอเล่นซ้ำเป็นล้านรอบ '
    'มันค่อยๆ จับสัญญาณได้เองว่า "อ้อ ตอน RSI ต่ำ + ADX สูง + แท่งกลับตัว... ควรซื้อ"'
)

add_analogy(doc,
    '🚗 เปรียบเทียบ: ขี่จักรยาน',
    'ตอนเด็กๆ คุณไม่ได้อ่านตำราฟิสิกส์เรื่องสมดุลก่อนขี่จักรยาน แต่ลองล้มไปเรื่อยๆ '
    'จนสมองคุณ "จำได้เอง" ว่าต้องเอนตัวยังไง บีบเบรกตอนไหน '
    'RL Bot ก็แบบเดียวกัน — มันล้ม (ขาดทุน) เป็นล้านครั้งในข้อมูลย้อนหลัง '
    'จนรู้วิธีทรงตัวเอง'
)

doc.add_page_break()

# ----------------- 3. รางวัล/ลงโทษ -----------------
add_heading(doc, '3. หัวใจสำคัญ: ระบบ "รางวัล" (Reward)', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))
add_para(doc,
    'นี่คือส่วนที่สำคัญที่สุด เพราะ "AI จะทำในสิ่งที่ได้รางวัลมากที่สุด" '
    'ถ้าเราออกแบบรางวัลผิด AI จะเรียนรู้ผิดทันที'
)

add_heading(doc, 'ตัวอย่างการ "ออกแบบรางวัลผิด"', level=2)
add_para(doc, 'V2: ให้รางวัลตามกำไรเฉยๆ', bold=True)
add_para(doc, '   → AI เรียนรู้ว่า "ไม่เทรดเลย ก็ไม่เสียอะไร" → ไม่เทรดเลยสักไม้!', color=RGBColor(0xCC, 0x00, 0x00))

add_para(doc, '\nV3: หักคะแนนตอนอยู่เฉย', bold=True)
add_para(doc, '   → AI เริ่มเทรด แต่เทรดมั่วเพราะกลัวโดนหัก → ขาดทุนรัวๆ', color=RGBColor(0xCC, 0x00, 0x00))

add_heading(doc, 'V4 - สูตรที่ใช้จริง (ใช้แล้วได้กำไร)', level=2)
v4 = [
    ['ปิดออเดอร์มีกำไร', '+ คะแนน × 50', 'รางวัลใหญ่เมื่อกำไรจริง'],
    ['กำไรเกิน 0.5%', '+ โบนัส 1.0', 'กระตุ้นให้กล้าถือไม้ดี'],
    ['ปิดออเดอร์ขาดทุน', '- คะแนน × 50', 'ลงโทษหนัก'],
    ['เปิดเทรด 1 ครั้ง', '- 0.005', 'หักนิดหน่อย กันเทรดมั่ว'],
]
add_table(doc, ['สถานการณ์', 'รางวัล/โทษ', 'เหตุผล'], v4)

add_callout(doc,
    '💡 บทเรียนสำคัญ',
    'การออกแบบ Reward คือ 80% ของความสำเร็จ! เหมือนคุณตั้งเป้าหมายให้ลูกน้อง — '
    'ถ้าตั้งผิด เขาจะวิ่งไปทางผิด เราเสียเวลาเทรน V2-V8 ก่อนจะเจอสูตร V4 ที่ถูกต้อง'
)

doc.add_page_break()

# ----------------- 4. workflow ภาพรวม -----------------
add_heading(doc, '4. ภาพรวมการทำงาน (7 ขั้นตอน)', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))
add_para(doc, 'นี่คือเส้นทางตั้งแต่ข้อมูลดิบ จนได้ EA ที่ใช้งานจริงใน MT5')

steps = [
    ['1', 'เก็บข้อมูล', 'ดึงราคาย้อนหลังจาก MT5 ออกมาเป็น CSV (ใช้ EA DataCollector)'],
    ['2', 'เตรียมข้อมูล', 'แปลงให้เป็นรูปแบบที่ AI กินได้ (normalize + แบ่ง train/test)'],
    ['3', 'เลือก Feature', 'เลือก indicator/pattern ที่มีผลกับราคาจริง'],
    ['4', 'ฝึก AI', 'ให้ AI เล่นซ้ำๆ เป็นล้านรอบจนเก่ง (PPO algorithm)'],
    ['5', 'ทดสอบย้อนหลัง', 'Backtest บนข้อมูลที่ AI ไม่เคยเห็น เพื่อเช็คว่าไม่ overfit'],
    ['6', 'แปลงเป็น EA', 'Export สมอง AI เป็น ONNX แล้วใส่ใน EA ของ MT5'],
    ['7', 'ใช้งานจริง', 'รัน EA บน demo/live account'],
]
add_table(doc, ['#', 'ขั้นตอน', 'ทำอะไร'], steps)

add_callout(doc,
    '🎯 ใช้เครื่องมือเดียวจบ',
    'ทุกขั้นตอนทำผ่านโปรแกรม "PyCaret Trainer GUI" ที่เป็นหน้าจอ Windows '
    'ปุ่มกดง่ายๆ ไม่ต้องพิมพ์โค้ด แค่ Import → Train → Backtest → Export '
    'แล้วเอาไฟล์ที่ได้ไปใส่ MT5'
)

doc.add_page_break()

# ----------------- 5. Phase by Phase ง่ายๆ -----------------
add_heading(doc, '5. แต่ละขั้นตอนทำอะไร (เจาะลึก)', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))

# 5.1
add_heading(doc, 'ขั้นที่ 1: เก็บข้อมูล', level=2)
add_para(doc,
    'เปิด MT5 → ใส่ EA "DataCollector" บนกราฟ EUR/USD H1 → กดรัน '
    'EA จะเขียนไฟล์ CSV ออกมา ที่มีข้อมูล:'
)
add_bullet(doc, 'OHLCV (เปิด/สูง/ต่ำ/ปิด/ปริมาณ)')
add_bullet(doc, 'Indicator: RSI, ATR, ADX, MACD, Bollinger Bands, Stochastic ฯลฯ')
add_bullet(doc, 'Multi-timeframe: ค่า RSI/Trend/ATR ของ Daily (ดูภาพใหญ่)')
add_bullet(doc, 'Candle Patterns 10 แบบ: Engulfing, Hammer, Star, Mat Hold ฯลฯ')

# 5.2
add_heading(doc, 'ขั้นที่ 2: เตรียมข้อมูล', level=2)
add_para(doc,
    'ราคา EUR/USD = 1.0850 แต่ RSI = 65 ตัวเลขขนาดต่างกัน 60 เท่า '
    'ถ้าโยนให้ AI ดิบๆ มันจะ "เห็น" ราคาเด่นกว่าทุกอย่าง '
    'เลยต้อง Normalize ให้ทุก feature อยู่ในสเกลเดียวกัน (เหมือนแปลงเป็น %) '
    'แล้วแบ่งข้อมูล:'
)
add_bullet(doc, 'Train set 80% — ใช้ฝึก AI')
add_bullet(doc, 'Test set 20% — เก็บไว้ทดสอบทีหลัง (AI ห้ามเห็น!)')

add_analogy(doc,
    '📚 เปรียบเทียบ: ติวข้อสอบ',
    'Train = หนังสือเรียน, Test = ข้อสอบจริง '
    'ถ้าให้นักเรียนเห็นข้อสอบก่อน ก็ไม่รู้ว่าเขาเก่งจริงไหม '
    'AI ก็เหมือนกัน ต้องเก็บ test set ไว้ห้ามเห็น'
)

# 5.3
add_heading(doc, 'ขั้นที่ 3: เลือก Feature ที่มีคุณภาพ', level=2)
add_para(doc,
    'ไม่ใช่ทุก indicator มีค่าเท่ากัน บางตัวเดาราคาได้แม่น บางตัวเป็น noise '
    'เราใช้เครื่องมือ "Feature Analysis" คำนวณว่า indicator ไหนสัมพันธ์กับราคาในอนาคตบ้าง '
)

corr = [
    ['d1_rsi (RSI ของ Daily)', '0.138', 'แรงสุด! ภาพใหญ่สำคัญมาก'],
    ['d1_trend', '0.095', 'เทรนด์ Daily'],
    ['rsi (H1)', '0.041', 'RSI ปกติ'],
    ['adx', '0.038', 'ความแรงเทรนด์'],
    ['macd_hist', '0.029', 'MACD'],
    ['candle_engulfing', '0.024', 'แท่ง Engulfing'],
]
add_table(doc, ['Feature', 'ค่า correlation', 'ความเห็น'], corr)

add_callout(doc,
    '💡 ค้นพบสำคัญ',
    'RSI ของ Daily Timeframe มีค่า correlation 0.138 = แรงกว่า indicator H1 ปกติ 3.4 เท่า! '
    'นี่คือเหตุผลที่เทรดเดอร์เก่งๆ มักดู Multi-timeframe — AI ก็ confirm สิ่งนี้'
)

doc.add_page_break()

# 5.4
add_heading(doc, 'ขั้นที่ 4: ฝึก AI (หัวใจของระบบ)', level=2)
add_para(doc,
    'ขั้นนี้คือที่ AI "เล่น" ข้อมูลย้อนหลังซ้ำๆ เป็นล้านรอบ '
    'ใช้ algorithm ชื่อ PPO (Proximal Policy Optimization) — ดังในวงการ AI'
)

add_heading(doc, 'AI เรียนยังไง (ทีละ episode)', level=3)
add_para(doc, '1. AI เริ่มที่จุดสุ่มในข้อมูล (วันไหนก็ได้)')
add_para(doc, '2. เห็นกราฟ 20 แท่งล่าสุด + indicator')
add_para(doc, '3. ตัดสินใจ: ซื้อ / ขาย / อยู่เฉย')
add_para(doc, '4. ระบบคำนวณ reward (กำไร/ขาดทุน × 50)')
add_para(doc, '5. AI ปรับสมองนิดนึงตาม reward')
add_para(doc, '6. ขยับไปแท่งถัดไป กลับไปขั้น 2')
add_para(doc, '7. ทำซ้ำเป็นล้านรอบ')

add_callout(doc,
    '⏱ ใช้เวลานานแค่ไหน?',
    '500,000 รอบ = ประมาณ 30 นาที - 1 ชม. (บนคอม CPU ปกติ)\n'
    '1,000,000 รอบ = 1-2 ชม.\n'
    '2,000,000 รอบ = 3-4 ชม. (สำหรับโมเดลใหญ่)'
)

# 5.5
add_heading(doc, 'ขั้นที่ 5: Backtest ตรวจฝีมือ', level=2)
add_para(doc,
    'พอ AI ฝึกเสร็จ เราเอามาทดสอบกับ Test set (ที่ AI ไม่เคยเห็น) '
    'ดูค่าสำคัญ:'
)
add_bullet(doc, 'Profit Factor (PF) — เป้าหมาย > 1.0 (กำไรรวมหารขาดทุนรวม)')
add_bullet(doc, 'Win Rate — % ของ trade ที่ชนะ')
add_bullet(doc, 'Max Drawdown — ขาดทุนสะสมสูงสุด')
add_bullet(doc, 'จำนวน Trade — กี่ไม้ในช่วงทดสอบ')

add_callout(doc,
    '🚨 Confidence Threshold (เคล็ดลับสำคัญ)',
    'เราตั้งให้ AI เทรดเฉพาะตอนที่ "มั่นใจสูง 95% ขึ้นไป" เท่านั้น '
    'พอใส่ filter นี้ จาก PF 0.95 (ขาดทุน) → กลายเป็น PF 1.03 (กำไร)! '
    'การเทรดน้อยแต่แม่น ดีกว่าเทรดเยอะแต่มั่ว'
)

# 5.6
add_heading(doc, 'ขั้นที่ 6: แปลงเป็น EA', level=2)
add_para(doc,
    'AI ที่ฝึกเสร็จเป็นไฟล์ Python (.zip) แต่ MT5 อ่านไม่เข้าใจ '
    'เราต้องแปลงเป็นไฟล์ ONNX (มาตรฐานกลางของ AI) ก่อน'
)

add_para(doc, 'กดปุ่ม "Export to ONNX" ในโปรแกรม จะได้ไฟล์ 3 ตัว:')
add_bullet(doc, 'model.onnx — สมอง AI (binary file)')
add_bullet(doc, 'model_config.mqh — ตั้งค่า (รายชื่อ feature, spread, confidence)')
add_bullet(doc, 'model_EA.mq5 — EA สำเร็จรูป พร้อม compile ใน MetaEditor')

add_analogy(doc,
    '📦 เปรียบเทียบ: บรรจุภัณฑ์',
    'AI ที่ฝึกใน Python เหมือนอาหารโฮมเมด → ต้องบรรจุใส่กล่อง (ONNX) '
    'เพื่อขายในซูเปอร์มาร์เก็ต (MT5) — รสชาติเหมือนเดิม แต่ห่อใหม่ให้เข้ากับร้าน'
)

# 5.7
add_heading(doc, 'ขั้นที่ 7: ใช้งานจริง', level=2)
add_para(doc, 'ก๊อปปี้ไฟล์ทั้งหมดเข้า MT5 → Compile ใน MetaEditor → Attach EA บนกราฟ EUR/USD H1')

add_callout(doc,
    '⚠ ทดลองก่อนเสมอ',
    'ห้ามเอา EA ใหม่ขึ้น Live Account ทันที! ทำตามลำดับนี้:\n'
    '1. Strategy Tester ใน MT5 (free, ไม่เสียเงิน)\n'
    '2. Demo Account (อย่างน้อย 2-4 สัปดาห์)\n'
    '3. Live Account ขนาดเล็กก่อน\n'
    '4. ค่อยเพิ่ม lot size ทีละนิด',
    bg='FFE0E0'
)

doc.add_page_break()

# ----------------- 6. Versions Journey -----------------
add_heading(doc, '6. การเดินทาง: V2 → V11', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))
add_para(doc,
    'เราไม่ได้สำเร็จในครั้งแรก! กว่าจะได้ Bot ที่ทำกำไร ต้องลองผิดลองถูก 11 versions:'
)
versions = [
    ['V2', 'Reward แบบง่าย', 'ขาดทุน', 'AI ไม่ยอมเทรดเลย'],
    ['V3', 'หักโทษตอนอยู่เฉย', 'ขาดทุน', 'AI เทรดมั่วซะ'],
    ['V4', 'สูตร reward ใหม่ ×50', 'เริ่มดี', 'พื้นฐานที่ใช้ต่อ'],
    ['V5-V8', 'ลอง feature ต่างๆ', 'เกือบ break-even', 'ใกล้แล้ว'],
    ['V9', 'เพิ่ม Daily timeframe', 'เกือบกำไร', 'ใกล้แล้วๆ'],
    ['V10', 'ใส่ confidence 95%', '✅ PF 1.03', 'ตัวแรกที่กำไร!'],
    ['V11', 'V10 + Candle Patterns', 'กำลังทดสอบ', 'หวังว่าดีกว่า V10'],
]
add_table(doc, ['Version', 'การเปลี่ยนแปลง', 'ผล', 'หมายเหตุ'], versions)

add_callout(doc,
    '💪 บทเรียน',
    'การฝึก AI ก็เหมือนการเทรน trader ใหม่ — ต้องลองผิด ปรับสูตร '
    'จนเจอ "วิธีคิด" ที่เหมาะกับตลาด แต่ละ version ใช้เวลา 1-3 ชม. ฝึก '
    'รวมแล้วใช้เวลาหลายสัปดาห์กว่าจะเจอตัวที่ดี'
)

doc.add_page_break()

# ----------------- 7. ข้อดี/ข้อเสีย -----------------
add_heading(doc, '7. ข้อดี vs ข้อเสีย ของ RL Bot', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))

add_heading(doc, '✅ ข้อดี', level=2)
add_bullet(doc, 'ไม่ต้องเขียนกฎเอง AI หากฎจาก data')
add_bullet(doc, 'จับ pattern ซับซ้อนที่คนคิดไม่ออก')
add_bullet(doc, 'พร้อมปรับเมื่อเจอข้อมูลใหม่ (retrain ได้)')
add_bullet(doc, 'ไม่มีอารมณ์ ไม่กลัว ไม่โลภ')
add_bullet(doc, 'ทำงาน 24/5 ไม่เหนื่อย')

add_heading(doc, '⚠ ข้อเสีย', level=2)
add_bullet(doc, 'ต้องการข้อมูลเยอะมาก (5-10 ปีขึ้นไป)')
add_bullet(doc, 'ฝึกใช้เวลานาน (1-4 ชม. ต่อ version)')
add_bullet(doc, '"Black box" — ไม่รู้ว่า AI คิดอะไรอยู่')
add_bullet(doc, 'ตลาดเปลี่ยนพฤติกรรม → ต้อง retrain เป็นระยะ')
add_bullet(doc, 'ไม่ใช่ "ปลูกเงิน" — ยังขาดทุนได้ถ้าเจอตลาดผิดสไตล์')

add_callout(doc,
    '🎯 RL Bot = เครื่องมือ ไม่ใช่เวทมนตร์',
    'มันเก่งกว่า EA แบบ if-else แต่ไม่ได้ "ทำเงินอัตโนมัติ" '
    'ยังต้องใช้คนคอยดู คอย retrain เมื่อตลาดเปลี่ยน '
    'มอง RL Bot เป็น "ลูกศิษย์เทรดเดอร์ที่ตัดสินใจเร็วและไม่มีอารมณ์" '
    'แต่ครู (เรา) ยังต้องดูแล',
    bg='E8F5E9'
)

doc.add_page_break()

# ----------------- 8. คำถามที่เจอบ่อย -----------------
add_heading(doc, '8. คำถามที่เจอบ่อย (FAQ)', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))

faq = [
    ['ต้องเขียนโค้ดเองไหม?',
     'ไม่ต้อง โปรแกรมมี GUI ปุ่มกดให้หมด แค่ Import → Train → Export'],
    ['ใช้กับคู่อื่นได้ไหม?',
     'ได้ ตอนเก็บข้อมูลในขั้นที่ 1 เปลี่ยนคู่ที่ MT5 + ฝึกใหม่ก็พอ '
     '(แต่ละคู่ต้อง model แยก)'],
    ['ใช้กับ TF อื่นได้ไหม?',
     'ได้ แค่เปลี่ยน TF ตอนเก็บข้อมูล (M15, H4, D1) แต่ TF เล็กเทรดเยอะ spread กิน'],
    ['ต้อง GPU ไหม?',
     'ไม่ต้อง CPU ปกติพอ (ฝึก 1-3 ชม.) มี GPU จะเร็วขึ้น 2-3 เท่า'],
    ['Bot ทำกำไร 100% เลยไหม?',
     'ไม่มี Bot ไหนทำกำไร 100% เป้าหมายคือ PF > 1.0 + Drawdown ที่รับได้'],
    ['ใช้กับ Crypto/หุ้นได้ไหม?',
     'ได้ ถ้าโบรกเกอร์รองรับใน MT5 + มีข้อมูลย้อนหลังพอ'],
    ['Retrain บ่อยแค่ไหน?',
     'แนะนำทุก 3-6 เดือน หรือเมื่อ performance เริ่มตก'],
    ['VPS ต้องสเปกแค่ไหน?',
     'EA รันใน MT5 ปกติ — VPS เกรดเริ่มต้น 2GB RAM พอ'],
]
add_table(doc, ['คำถาม', 'คำตอบ'], faq)

doc.add_page_break()

# ----------------- 9. คำศัพท์ -----------------
add_heading(doc, '9. คำศัพท์ที่ควรรู้', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))

terms = [
    ['RL', 'Reinforcement Learning — เรียนรู้จาก reward'],
    ['Agent', 'AI ผู้เล่นที่เรากำลังฝึก'],
    ['Episode', '1 รอบเทรด (เริ่มจุดสุ่ม → จบเมื่อหมด data)'],
    ['Reward', 'คะแนน + หรือ - ที่ AI ได้รับ'],
    ['Feature', 'ข้อมูลที่ AI ใช้ตัดสินใจ (RSI, ATR, ฯลฯ)'],
    ['PPO', 'Algorithm หลักที่ใช้ฝึก (มาตรฐานวงการ)'],
    ['ONNX', 'มาตรฐานไฟล์ AI ที่ MT5 อ่านได้'],
    ['Overfit', 'AI จำข้อมูลเก่าได้ดี แต่เทรดข้อมูลใหม่ไม่ได้'],
    ['Profit Factor', 'กำไรรวม ÷ ขาดทุนรวม > 1.0 = กำไร'],
    ['Drawdown', 'ขาดทุนสะสมสูงสุดจากจุดสูงสุด'],
    ['Confidence', 'ความมั่นใจของ AI (0-100%)'],
    ['Backtest', 'ทดสอบกับข้อมูลย้อนหลัง'],
    ['Walk-forward', 'Backtest แบบขยับหน้าต่างเรื่อยๆ ใกล้ live เทรด'],
    ['Normalize', 'ปรับสเกลข้อมูลให้อยู่ในระดับเดียวกัน'],
]
add_table(doc, ['คำศัพท์', 'ความหมาย'], terms)

doc.add_page_break()

# ----------------- 10. สรุป -----------------
add_heading(doc, '10. สรุปสำหรับเทรดเดอร์', level=1, color=RGBColor(0x2E, 0x5B, 0xBA))

add_para(doc, 'RL Bot คืออะไร?', bold=True, size=14)
add_para(doc,
    'AI ที่เรียนการเทรดจากการลองผิดลองถูกซ้ำๆ เป็นล้านรอบในข้อมูลย้อนหลัง '
    'จนหากฎการตัดสินใจของตัวเอง — โดยไม่ต้องมีใครสอน'
)

add_para(doc, '\nต่างจาก EA ปกติยังไง?', bold=True, size=14)
add_para(doc,
    'EA ปกติใช้กฎที่คนเขียนตายตัว (if RSI<30 ซื้อ) '
    'RL Bot ใช้กฎที่ AI หาเองจาก data — ปรับตัวได้ดีกว่า แต่ก็ "อ่านยาก" กว่า'
)

add_para(doc, '\nผลลัพธ์ที่ได้?', bold=True, size=14)
add_para(doc,
    'EA ตัวหนึ่งสำหรับ MT5 ที่:'
)
add_bullet(doc, 'ใช้งานเหมือน EA ปกติ (Attach บนกราฟ → รัน)')
add_bullet(doc, 'มี ONNX model ฝังไว้ข้างใน — ตัดสินใจ buy/sell/hold')
add_bullet(doc, 'มี Confidence Filter — เทรดเฉพาะตอนมั่นใจ')
add_bullet(doc, 'มี Session Filter — ไม่เทรดตอนตลาดปิด')
add_bullet(doc, 'รองรับ Strategy Tester ของ MT5')

add_para(doc, '\nใครเหมาะใช้?', bold=True, size=14)
add_bullet(doc, 'เทรดเดอร์ที่อยากลองเครื่องมือใหม่ๆ')
add_bullet(doc, 'คนที่อยาก automate แต่ไม่อยากเขียนกฎเอง')
add_bullet(doc, 'คนที่มีข้อมูลย้อนหลังดีๆ เยอะๆ')
add_bullet(doc, 'คนที่พร้อมศึกษา ลองผิดลองถูก ไม่ใจร้อน')

add_para(doc, '\nใครไม่เหมาะ?', bold=True, size=14)
add_bullet(doc, 'คนที่อยาก "EA รวยเร็ว" — ไม่มีอยู่จริง')
add_bullet(doc, 'คนที่ไม่ชอบ retrain เมื่อตลาดเปลี่ยน')
add_bullet(doc, 'คนไม่มีเวลาดู demo อย่างน้อย 2-4 สัปดาห์')

doc.add_paragraph('\n\n')
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = end.add_run('— จบเอกสาร —')
run.italic = True
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

end2 = doc.add_paragraph()
end2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = end2.add_run('\nขอให้สนุกกับการเทรน Bot ของคุณ! 🚀')
run.bold = True
run.font.color.rgb = RGBColor(0x2E, 0x5B, 0xBA)

# ----------------- SAVE -----------------
output_path = r'C:\Users\omesb\Documents\claude code\pycaret_trainer\RL_Trading_Simple.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
